import base64
import hashlib
import inspect
import json
import os
import sys
import tempfile
import threading
import time
import tracemalloc
import unittest
from datetime import datetime
from pathlib import Path
from unittest import mock

import fitz
from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image

from api.api import API
from api.core.context import TaskCancelled, TaskContext, run_process, task_context
from api.core.outputs import atomic_output
from api.core.store import StateStore
from pyapp.config.config import Config


class DesktopBridgeTests(unittest.TestCase):
    def setUp(self):
        self.temporary = tempfile.TemporaryDirectory()
        self.root = Path(self.temporary.name)
        self.patches = [mock.patch.object(Config, name, str(self.root)) for name in ("appDataDir", "downloadDir")]
        for patch in self.patches:
            patch.start()
        self.api = API()

    def tearDown(self):
        self.api.task_shutdown()
        self.api.workflow_stop()
        for patch in self.patches:
            patch.stop()
        self.temporary.cleanup()

    def wait_task(self, identity, timeout=30):
        deadline = time.monotonic() + timeout
        while time.monotonic() < deadline:
            task = self.api.task_get({"id": identity})["task"]
            if task["status"] in {"success", "partial", "failed", "canceled", "interrupted"}:
                return task
            time.sleep(0.03)
        self.fail("任务没有在时限内结束")

    def test_bridge_word_outputs_docx_and_never_replaces_input(self):
        paths = []
        for text in ("first document", "second document"):
            path = self.root / (text + ".docx")
            doc = Document()
            doc.add_paragraph(text)
            doc.save(path)
            paths.append(str(path))
        before = hashlib.sha256(Path(paths[0]).read_bytes()).hexdigest()
        result = self.api.word_merge({"files": paths})
        self.assertEqual(result["code"], 0, result)
        self.assertEqual(Path(result["output"]).suffix, ".docx")
        self.assertIn("second document", "\n".join(p.text for p in Document(result["output"]).paragraphs))
        again = self.api.word_merge({"files": paths, "outputPath": paths[0]})
        self.assertEqual(again["code"], 0, again)
        self.assertNotEqual(again["output"], paths[0])
        self.assertEqual(hashlib.sha256(Path(paths[0]).read_bytes()).hexdigest(), before)
        directory_only = self.api.word_merge({'files': paths, 'outputDir': str(self.root / 'chosen-output')})
        self.assertEqual(directory_only['code'], 0, directory_only)
        self.assertEqual(Path(directory_only['output']).parent, self.root / 'chosen-output')
        self.assertEqual(Path(directory_only['output']).suffix, '.docx')
        self.assertEqual(inspect.getfullargspec(self.api.word_merge).args, ["self", "options"])
        self.assertFalse(any("mindmap" in name.lower() for name in dir(self.api)))

    def test_bridge_seal_preview_matches_export_with_texture(self):
        fonts = [
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/System/Library/Fonts/Helvetica.ttc"),
        ]
        font = next((path for path in fonts if path.exists()), None)
        if font is None:
            self.skipTest("当前环境没有测试字体")
        texture = self.root / "texture.png"
        Image.effect_noise((1024, 1024), 40).save(texture)
        options = {
            "template": "ellipse",
            "topText": "PPX",
            "middleText": "TEST",
            "bottomText": "123",
            "fontChinesePath": str(font),
            "fontLatinPath": str(font),
            "texturePath": str(texture),
            "outputDir": str(self.root),
        }
        preview = self.api.seal_generate_preview(options)
        exported = self.api.seal_generate({**options, "mode": "export"})
        self.assertEqual(exported["code"], 0, exported)
        self.assertEqual(preview["preview"], exported["preview"])
        self.assertEqual(base64.b64decode(preview["preview"].split(",", 1)[1]), Path(exported["output"]).read_bytes())

    def test_excel_numbers_dates_leading_zeroes_and_formulas(self):
        source = self.root / "typed.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.append(["n", "code", "date", "double"])
        for index, value in enumerate([10, 2, 1], 2):
            ws.append([value, "0012", datetime(2026, 9, value), f"=A{index}*2"])
            ws.cell(index, 1).number_format = "0000"
        wb.save(source)
        wb.close()
        result = self.api.excel_process(
            {"filePath": str(source), "sortBy": "n", "exportCombined": True, "exportGroups": False, "exportJson": False}
        )
        self.assertEqual(result["code"], 0, result)
        output = load_workbook(result["combinedPath"])
        try:
            rows = list(output.active.iter_rows(min_row=2))
            self.assertEqual([row[0].value for row in rows], [1, 2, 10])
            self.assertTrue(all(row[1].value == "0012" for row in rows))
            self.assertTrue(all(isinstance(row[2].value, datetime) for row in rows))
            self.assertEqual([row[3].value for row in rows], ["=A2*2", "=A3*2", "=A4*2"])
            self.assertEqual(rows[0][0].number_format, "0000")
        finally:
            output.close()
        cached = self.api.excel_preview({"filePath": str(source), "formulaPolicy": "values"})
        self.assertNotEqual(cached["code"], 0)
        self.assertIn("缓存", cached["msg"])

    def test_excel_merge_uses_header_mapping(self):
        tables = []
        for name, header, row in [("a", ["name", "count"], ["a", 10]), ("b", ["数量", "名称"], [2, "b"])]:
            path = self.root / f"{name}.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.append(header)
            ws.append(row)
            wb.save(path)
            wb.close()
            tables.append({"path": str(path), "fieldMapping": {"name": "名称", "count": "数量"} if name == "b" else {}})
        result = self.api.excel_merge_tables({"tables": tables, "outputDir": str(self.root)})
        self.assertEqual(result["code"], 0, result)
        wb = load_workbook(result["output"])
        try:
            self.assertEqual(list(wb.active.values), [("name", "count"), ("a", 10), ("b", 2)])
        finally:
            wb.close()

    def test_pdf_thousand_pages_preview_and_reorder(self):
        source = self.root / "large.pdf"
        doc = fitz.open()
        for number in range(1000):
            page = doc.new_page(width=100, height=100)
            page.insert_text((10, 30), f"page {number + 1}")
        doc.save(source)
        doc.close()
        preview = self.api.pdf_page_preview({"filePath": str(source), "offset": 976, "limit": 24, "width": 100})
        self.assertEqual(preview["code"], 0, preview)
        self.assertEqual(preview["pageCount"], 1000)
        self.assertEqual([page["page"] for page in preview["pages"]], list(range(977, 1001)))
        result = self.api.pdf_page_workbench(
            {"filePath": str(source), "pageOrder": [1000, 1, 500], "sourceSignature": preview["sourceSignature"]}
        )
        self.assertEqual(result["code"], 0, result)
        with fitz.open(result["output"]) as output:
            self.assertEqual([page.get_text().strip() for page in output], ["page 1000", "page 1", "page 500"])
        empty = self.api.pdf_page_workbench({"filePath": str(source), "pageOrder": []})
        self.assertNotEqual(empty["code"], 0)

    @mock.patch.dict(os.environ, {'PYTHONIOENCODING': 'cp1252', 'PYTHONUTF8': '0'})
    def test_partial_batch_only_retries_failed_input(self):
        good = self.root / "图像.png"
        missing = self.root / "重试.png"
        Image.new("RGBA", (80, 40), (50, 120, 200, 80)).save(good)
        submitted = self.api.task_submit(
            {
                "method": "image_rotate_flip",
                "args": [
                    {"files": [str(good), str(missing)], "operation": "rotate90", "outputDir": str(self.root / "out")}
                ],
            }
        )
        task = self.wait_task(submitted["taskId"])
        self.assertEqual(task["status"], "partial", task)
        self.assertEqual(len(task["outputs"]), 1)
        with Image.open(task["outputs"][0]["path"]) as image:
            self.assertEqual(image.size, (40, 80))
            self.assertEqual(image.mode, "RGBA")
        Image.new("RGBA", (20, 10)).save(missing)
        retried = self.api.task_retry({"id": task["id"]})
        self.assertEqual(retried["code"], 0, retried)
        finished = self.wait_task(retried["taskId"])
        self.assertEqual(finished["status"], "success", finished)
        self.assertEqual(finished["args"][0]["files"], [str(missing)])
        self.assertEqual(len(list((self.root / "out").glob("*.png"))), 2)

    def test_canceled_batch_keeps_completed_outputs(self):
        source = self.root / "image.png"
        Image.new("RGBA", (1000, 1000), (30, 40, 60, 50)).save(source)
        submitted = self.api.task_submit(
            {
                "method": "image_rotate_flip",
                "args": [{"files": [str(source)] * 300, "outputDir": str(self.root / "out")}],
            }
        )
        identity = submitted["taskId"]
        deadline = time.monotonic() + 30
        while time.monotonic() < deadline:
            task = self.api.task_get({"id": identity})["task"]
            if task["outputs"]:
                break
            time.sleep(0.03)
        self.assertTrue(task["outputs"], task)
        canceled = self.api.task_cancel({"id": identity})
        self.assertEqual(canceled["task"]["status"], "canceling")
        finished = self.wait_task(identity, timeout=8)
        self.assertEqual(finished["status"], "canceled", finished)
        self.assertTrue(all(Path(asset["path"]).is_file() for asset in finished["outputs"]))
        self.assertFalse(list((self.root / "out").glob(".*.tmp*")))


class PersistenceAndProcessTests(unittest.TestCase):
    def test_json_migration_is_repeatable_and_preserves_backup(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            legacy = root / "old.json"
            legacy.write_text(json.dumps({"tasks": [{"id": "old", "status": "success"}]}), encoding="utf-8")
            store = StateStore(root)
            first = store.load("tasks", legacy, {})
            store.save("tasks", {**first, "tasks": first["tasks"] + [{"id": "new"}]})
            second = StateStore(root).load("tasks", legacy, {})
            self.assertEqual([item["id"] for item in second["tasks"]], ["old", "new"])
            self.assertEqual(legacy.read_bytes(), legacy.with_name("old.json.pre-sqlite.bak").read_bytes())

    def test_atomic_failure_leaves_input_untouched(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "output.txt"
            path.write_text("original")
            with self.assertRaises(RuntimeError):
                with atomic_output(path) as (temporary, _):
                    temporary.write_text("incomplete")
                    raise RuntimeError("disk write failed")
            self.assertEqual(path.read_text(), "original")
            self.assertEqual(list(Path(directory).iterdir()), [path])

    def test_blocking_process_is_terminated_on_cancel(self):
        context = TaskContext()
        timer = threading.Timer(0.25, context.cancel.set)
        timer.start()
        started = time.monotonic()
        try:
            with task_context(context), self.assertRaises(TaskCancelled):
                run_process([sys.executable, "-c", "import time; time.sleep(30)"], capture_output=True)
        finally:
            timer.cancel()
            timer.join()
        self.assertLess(time.monotonic() - started, 5)

    def test_preview_of_hundred_thousand_rows_is_bounded(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "large.xlsx"
            wb = Workbook(write_only=True)
            ws = wb.create_sheet()
            ws.append(["number", "code"])
            for number in range(100000):
                ws.append([number, f"{number:08}"])
            wb.save(path)
            wb.close()
            api = API()
            tracemalloc.start()
            try:
                result = api.excel_preview({"filePath": str(path), "limit": 10})
                _, peak = tracemalloc.get_traced_memory()
            finally:
                tracemalloc.stop()
            self.assertEqual(result["code"], 0, result)
            self.assertEqual(len(result["sample"]), 10)
            self.assertEqual(result["sample"][0]["code"], "00000000")
            self.assertLess(peak, 64 * 1024 * 1024)


if __name__ == "__main__":
    unittest.main()
