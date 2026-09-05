#!/usr/bin/env python3
# -*- coding: utf-8 -*-
'''
Author: Codex
Date: 2026-06-07
Description: Word(.docx) 工具相关 API —— 拆分(多文件)、切割(保留指定页)、合并

“按页码”相关能力说明：
  .docx 没有原生的“页”概念（页由排版引擎实时计算）。本模块通过调用本机
  LibreOffice 将文档转换为 PDF，得到与 Word 基本一致的真实分页，再用 PyMuPDF
  解析每个内容块所落的页码，建立“内容块 → 页码”映射。真正的拆分/切割始终在
  原始文档上以“复制后删除多余元素”的方式完成，尽可能保留原文档格式；跨页内容块在边界不完整时拒绝裁切。
'''

import base64
import copy
import re
import shutil
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from shutil import which
from typing import Dict, Iterable, List, Tuple

import fitz  # PyMuPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from api.core.context import checkpoint, iter_progress, run_process
from api.core.outputs import atomic_output, write_output
from api.utils.validators import ensure_output_directory

_MARKER_RE = re.compile(r'PPXMK(\d+)Z')


class WordTool():
    '''Word(.docx) 相关功能：拆分、切割、合并'''

    def _ensure_word_file(self, file_path: str) -> Path:
        if not file_path:
            raise ValueError('请选择 Word 文件')
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f'文件不存在：{path}')
        if path.suffix.lower() == '.doc':
            raise ValueError('暂不支持旧版 .doc 格式，请先另存为 .docx')
        if path.suffix.lower() != '.docx':
            raise ValueError('仅支持 .docx 文件')
        return path

    def _timestamp(self) -> str:
        return datetime.now().strftime('%Y%m%d_%H%M%S')

    def _compose_output_path(self, directory: str, filename: str, source: Path, suffix: str) -> str:
        if directory or filename:
            safe_dir = Path(directory) if directory else source.parent
            safe_dir.mkdir(parents=True, exist_ok=True)
            filename = filename or f'{source.stem}_{suffix}_{self._timestamp()}.docx'
            if Path(filename).name != filename:
                raise ValueError('输出文件名不能包含目录分隔符')
            safe_name = filename if filename.lower().endswith('.docx') else f'{filename}.docx'
            return str(safe_dir / safe_name)
        return ''

    def _resolve_output_path(self, source: Path, output_path: str, suffix: str) -> Path:
        if output_path:
            dest = Path(output_path)
            dest.parent.mkdir(parents=True, exist_ok=True)
        else:
            dest = source.parent / f'{source.stem}_{suffix}_{self._timestamp()}.docx'
        return dest

    def _parse_page_spec(self, spec: str, total: int) -> List[int]:
        '''解析 "1-3,5,8" 形式的页码表达式，返回去重后的有序页码列表。'''
        pages: List[int] = []
        if not spec:
            return pages
        for chunk in str(spec).replace('，', ',').split(','):
            chunk = chunk.strip()
            if not chunk:
                continue
            try:
                if '-' in chunk:
                    start_str, end_str = chunk.split('-', 1)
                    start, end = int(start_str.strip()), int(end_str.strip())
                    if start > end:
                        start, end = end, start
                    for page in range(start, end + 1):
                        if 1 <= page <= total:
                            pages.append(page)
                else:
                    page = int(chunk)
                    if 1 <= page <= total:
                        pages.append(page)
            except (ValueError, TypeError):
                continue
        seen, unique = set(), []
        for page in pages:
            if page not in seen:
                seen.add(page)
                unique.append(page)
        return unique

    # --- 文档结构辅助 -------------------------------------------------

    def _content_blocks(self, document: Document) -> List:
        '''按文档顺序返回正文中的段落(w:p)与表格(w:tbl)元素，忽略结尾的 sectPr。'''
        blocks = []
        for child in document.element.body.iterchildren():
            if child.tag == qn('w:p'):
                blocks.append(child)
            elif child.tag == qn('w:tbl'):
                blocks.extend(child.findall(qn('w:tr')))
        return blocks

    def _has_page_break(self, paragraph_el) -> bool:
        '''段落是否包含手动分页符（run 级 <w:br w:type="page"/>）。'''
        for br in paragraph_el.iter(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
        return False

    def _has_section_break(self, paragraph_el) -> bool:
        '''段落是否为分节符（pPr 内嵌 sectPr，意味着下一段开始新节/新页）。'''
        pPr = paragraph_el.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            return True
        return False

    def _style_id_of(self, paragraph_el) -> str:
        pPr = paragraph_el.find(qn('w:pPr'))
        if pPr is None:
            return ''
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            return ''
        return pStyle.get(qn('w:val')) or ''

    def _is_heading(self, paragraph_el, id_to_name: Dict[str, str], level: int) -> bool:
        if paragraph_el.tag != qn('w:p'):
            return False
        style_id = self._style_id_of(paragraph_el)
        if not style_id:
            return False
        name = (id_to_name.get(style_id) or '').strip().lower()
        sid = style_id.strip().lower()
        if name in {f'heading {level}', f'标题 {level}'}:
            return True
        if sid == f'heading{level}':
            return True
        return False

    def _segments_by_count(self, total: int, per_file: int) -> List[range]:
        per_file = max(1, per_file)
        return [range(start, min(start + per_file, total)) for start in range(0, total, per_file)]

    def _segments_by_boundaries(self, total: int, boundaries: List[int]) -> List[range]:
        '''根据切点(每段起始下标)生成连续区间。'''
        points = sorted({b for b in boundaries if 0 <= b < total})
        if not points or points[0] != 0:
            points = [0] + points
        segments: List[range] = []
        for idx, start in enumerate(points):
            end = points[idx + 1] if idx + 1 < len(points) else total
            if start < end:
                segments.append(range(start, end))
        return segments

    @staticmethod
    def _section_properties(blocks, body):
        boundaries = []
        inherited = {}
        for index, block in enumerate(blocks):
            section = block.find('./' + qn('w:pPr') + '/' + qn('w:sectPr')) if block.tag == qn('w:p') else None
            if section is not None:
                boundaries.append((index, section))
        final = body.find(qn('w:sectPr'))
        if final is not None:
            boundaries.append((len(blocks) - 1, final))
        mapping, start = {}, 0
        for section_index, (end, properties) in enumerate(boundaries):
            properties = copy.deepcopy(properties)
            for reference in list(properties):
                if reference.tag in {qn('w:headerReference'), qn('w:footerReference')}:
                    inherited[(reference.tag, reference.get(qn('w:type')))] = copy.deepcopy(reference)
            explicit = {(reference.tag, reference.get(qn('w:type'))) for reference in properties}
            for key, reference in inherited.items():
                if key not in explicit:
                    properties.insert(0, copy.deepcopy(reference))
            for index in range(start, end + 1):
                mapping[index] = (section_index, properties)
            start = end + 1
        return mapping

    def _write_segment(self, source: Path, dest: Path, keep_indices) -> Path:
        '''复制源文件后裁剪正文，仅保留指定下标的内容块，最大限度保留样式/图片/页眉页脚。'''
        checkpoint()
        with atomic_output(dest) as (temporary, final):
            shutil.copyfile(source, temporary)
            doc = Document(str(temporary))
            blocks = self._content_blocks(doc)
            keep = set(keep_indices)
            body = doc.element.body
            section_map = self._section_properties(blocks, body)
            for table in body.findall(qn('w:tbl')):
                rows = table.findall(qn('w:tr'))
                if any(index in keep and block in rows for index, block in enumerate(blocks)):
                    for index, block in enumerate(blocks):
                        if block in rows and block.find('./' + qn('w:trPr') + '/' + qn('w:tblHeader')) is not None:
                            keep.add(index)
            for index, block in enumerate(blocks):
                if index not in keep:
                    block.getparent().remove(block)
            for table in list(body.findall(qn('w:tbl'))):
                if not table.findall(qn('w:tr')):
                    body.remove(table)
            retained = [(index, block) for index, block in enumerate(blocks) if index in keep]
            for _, block in retained:
                if block.tag == qn('w:p'):
                    section = block.find('./' + qn('w:pPr') + '/' + qn('w:sectPr'))
                    if section is not None:
                        section.getparent().remove(section)
            existing = body.find(qn('w:sectPr'))
            if existing is not None:
                body.remove(existing)
            for position, (index, block) in enumerate(retained):
                section_index, properties = section_map[index]
                is_last = position == len(retained) - 1
                if not is_last and section_map[retained[position + 1][0]][0] == section_index:
                    continue
                if is_last:
                    body.append(copy.deepcopy(properties))
                else:
                    if block.tag == qn('w:tr'):
                        paragraph = OxmlElement('w:p')
                        block.getparent().addnext(paragraph)
                    else:
                        paragraph = block
                    p_pr = paragraph.find(qn('w:pPr'))
                    if p_pr is None:
                        p_pr = OxmlElement('w:pPr')
                        paragraph.insert(0, p_pr)
                    p_pr.append(copy.deepcopy(properties))
            doc.save(str(temporary))
        return final

    # --- 真实分页(LibreOffice + PyMuPDF) ------------------------------

    def _locate_soffice(self) -> str:
        '''定位本机 LibreOffice 可执行文件。'''
        candidates: List[str] = []
        if sys.platform == 'darwin':
            candidates.append('/Applications/LibreOffice.app/Contents/MacOS/soffice')
        elif sys.platform.startswith('win'):
            candidates += [
                r'C:\Program Files\LibreOffice\program\soffice.exe',
                r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            ]
        for name in ('soffice', 'libreoffice'):
            found = which(name)
            if found:
                candidates.append(found)
        for candidate in candidates:
            if candidate and Path(candidate).exists():
                return candidate
        raise RuntimeError('未检测到 LibreOffice，按页码功能需要先安装 LibreOffice')

    def _insert_marker(self, block, index: int, at_end=False) -> None:
        '''在内容块起始处插入一个 1pt 的唯一标记 run，仅用于分页定位。'''
        if block.tag == qn('w:tr'):
            for cell in block.findall(qn('w:tc')):
                paragraphs = cell.findall('.//' + qn('w:p'))
                if paragraphs:
                    self._insert_marker(paragraphs[-1] if at_end else paragraphs[0], index, at_end)
            return
        if block.tag == qn('w:p'):
            target_p = block
        elif block.tag in (qn('w:tbl'), qn('w:tr')):
            paragraphs = block.findall('.//' + qn('w:p'))
            target_p = (paragraphs[-1] if at_end else paragraphs[0]) if paragraphs else None
            if target_p is None:
                return
        else:
            return
        token = f'PPXMK{index:05d}Z'
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        for tag in ('w:sz', 'w:szCs'):
            el = OxmlElement(tag)
            el.set(qn('w:val'), '2')  # 半磅，2 => 1pt，尽量减少对排版的扰动
            rPr.append(el)
        run.append(rPr)
        text = OxmlElement('w:t')
        text.set(qn('xml:space'), 'preserve')
        text.text = token
        run.append(text)
        if at_end:
            target_p.append(run)
            return
        pPr = target_p.find(qn('w:pPr'))
        if pPr is not None:
            pPr.addnext(run)
        else:
            target_p.insert(0, run)

    def _convert_to_pdf(self, soffice: str, src_docx: Path, out_dir: Path) -> Path:
        out_dir.mkdir(parents=True, exist_ok=True)
        profile = out_dir / 'lo_profile'
        cmd = [
            soffice, '--headless', '--norestore', '--invisible', '--nologo',
            f'-env:UserInstallation={profile.as_uri()}',
            '--convert-to', 'pdf', '--outdir', str(out_dir), str(src_docx)
        ]
        try:
            run_process(cmd, check=True, capture_output=True, text=False, timeout=180)
        except subprocess.TimeoutExpired:
            raise RuntimeError('LibreOffice 转换超时')
        except subprocess.CalledProcessError as exc:
            detail = (exc.stderr or b'').decode('utf-8', 'ignore').strip()
            raise RuntimeError(f'LibreOffice 转换失败：{detail or exc.returncode}')
        pdf = out_dir / f'{src_docx.stem}.pdf'
        if not pdf.exists():
            raise RuntimeError('LibreOffice 未生成 PDF')
        return pdf

    def _paginate(self, source):
        spans, total = self._page_spans(source)
        return [start for start, _ in spans], total

    def _page_spans(self, source):
        soffice = self._locate_soffice()
        with tempfile.TemporaryDirectory() as td:
            tmp = Path(td)
            marked = tmp / 'marked.docx'
            doc = Document(str(source))
            blocks = self._content_blocks(doc)
            control_only = {index for index, block in enumerate(blocks)
                            if block.tag == qn('w:p') and (self._has_page_break(block) or self._has_section_break(block))
                            and not any(node.text for node in block.iter(qn('w:t')))
                            and not list(block.iter(qn('w:drawing')))}
            for index, block in enumerate(blocks):
                self._insert_marker(block, index * 2)
                self._insert_marker(block, index * 2 + 1, at_end=True)
            doc.save(str(marked))
            pdf = self._convert_to_pdf(soffice, marked, tmp)
            original_pdf = self._convert_to_pdf(soffice, source, tmp / 'original')
            token_page = {}
            with fitz.open(str(pdf)) as pdoc, fitz.open(str(original_pdf)) as original:
                total_pages = pdoc.page_count
                if total_pages != original.page_count:
                    raise ValueError('定位标记影响分页，请改用结构拆分或 PDF 页面工具')
                for pno in range(total_pages):
                    checkpoint()
                    text = pdoc[pno].get_text('text')
                    normalized = re.sub(r'\s+', '', _MARKER_RE.sub('', text))
                    if normalized != re.sub(r'\s+', '', original[pno].get_text('text')):
                        raise ValueError(f'第 {pno + 1} 页无法可靠定位，请改用结构拆分或 PDF 页面工具')
                    for match in _MARKER_RE.finditer(text):
                        token = int(match.group(1))
                        block = blocks[token // 2]
                        repeated_header = block.tag == qn('w:tr') and block.find('./' + qn('w:trPr') + '/' + qn('w:tblHeader')) is not None
                        if token % 2 and not repeated_header:
                            token_page[token] = max(token_page.get(token, 1), pno + 1)
                        else:
                            token_page.setdefault(token, pno + 1)
            spans = []
            for index in range(len(blocks)):
                if index in control_only:
                    spans.append((0, 0))
                    continue
                if index * 2 not in token_page or index * 2 + 1 not in token_page:
                    raise ValueError(f'无法定位第 {index + 1} 个内容块，请使用结构拆分')
                spans.append((token_page[index * 2], token_page[index * 2 + 1]))
            return spans, max(total_pages, 1)

    def _indices_for_pages(self, spans, selected):
        keep = []
        for index, (start, end) in enumerate(spans):
            occupied = set(range(start, end + 1))
            if occupied.intersection(selected):
                if not occupied.issubset(selected):
                    raise ValueError(f'第 {index + 1} 个段落或表格行跨越 {start}–{end} 页；请包含完整内容块，或用 PDF 页面工具精确裁切')
                keep.append(index)
        return keep

    def word_preview(self, options=None):
        try:
            opts = options or {}
            source = self._ensure_word_file(opts.get('filePath', ''))
            doc = Document(str(source))
            blocks = self._content_blocks(doc)
            offset, limit = max(0, int(opts.get('offset') or 0)), min(50, max(1, int(opts.get('limit') or 20)))
            outline = [{'index': index, 'kind': 'tableRow' if block.tag == qn('w:tr') else 'paragraph',
                        'text': ''.join(node.text or '' for node in block.iter(qn('w:t')))[:500], 'style': self._style_id_of(block),
                        'sectionBreak': self._has_section_break(block)}
                       for index, block in enumerate(blocks[offset:offset + limit], offset)]
            result = {'code': 0, 'msg': '结构预览完成', 'blocks': outline, 'blockCount': len(blocks),
                      'sectionCount': len(doc.sections), 'pages': []}
            if opts.get('renderPages'):
                with tempfile.TemporaryDirectory() as temp:
                    pdf = self._convert_to_pdf(self._locate_soffice(), source, Path(temp))
                    with fitz.open(pdf) as document:
                        result['pageCount'] = document.page_count
                        page_start = max(0, int(opts.get('pageOffset') or 0))
                        for index in range(page_start, min(page_start + 6, document.page_count)):
                            pixmap = document[index].get_pixmap(matrix=fitz.Matrix(0.8, 0.8), alpha=False)
                            result['pages'].append({'page': index + 1, 'preview': 'data:image/png;base64,' + base64.b64encode(pixmap.tobytes('png')).decode()})
            return result
        except Exception as exc:
            return {'code': -1, 'msg': f'预览失败：{exc}'}

    # --- 对外 API ----------------------------------------------------

    def word_page_count(self, options: Dict = None):
        '''返回文档的真实页数（供前端按页码功能展示/校验）。'''
        try:
            opts = options if isinstance(options, dict) else {}
            source = self._ensure_word_file(opts.get('filePath', ''))
            _, total_pages = self._paginate(source)
            return {'code': 0, 'msg': f'共 {total_pages} 页', 'pages': total_pages}
        except Exception as exc:
            return {'code': -1, 'msg': f'页数计算失败：{exc}'}

    def word_split(self, options: Dict = None):
        '''Word 文档拆分为多个 .docx（保留格式）

        支持的模式（mode）：
          - pages      : 按真实页码，每 N 页一个文件（依赖 LibreOffice）
          - paragraphs : 每 N 个内容块（段落/表格）一个文件
          - pagebreak  : 在手动分页符 / 分节符处拆分
          - heading    : 在指定级别的标题处拆分
        '''
        try:
            opts = options if isinstance(options, dict) else {}
            source = self._ensure_word_file(opts.get('filePath', ''))
            mode = str(opts.get('mode') or 'pages').lower()
            per_file = int(opts.get('paragraphsPerFile') or 10)
            pages_per_file = max(1, int(opts.get('pagesPerFile') or 1))
            heading_level = int(opts.get('headingLevel') or 1)

            out_dir = ensure_output_directory(source, opts.get('outputDir', ''), 'split')

            doc = Document(str(source))
            blocks = self._content_blocks(doc)
            total = len(blocks)
            if total == 0:
                raise ValueError('文档没有可拆分的内容')

            extra_msg = ''
            if mode == 'pages':
                spans, total_pages = self._page_spans(source)
                segments = []
                start_page = 1
                while start_page <= total_pages:
                    end_page = min(start_page + pages_per_file - 1, total_pages)
                    idxs = self._indices_for_pages(spans, set(range(start_page, end_page + 1)))
                    if idxs:
                        segments.append(idxs)
                    start_page = end_page + 1
                if not segments:
                    raise ValueError('未能按页码拆分')
                extra_msg = f'（共 {total_pages} 页）'
            elif mode == 'paragraphs':
                segments = self._segments_by_count(total, per_file)
            elif mode == 'pagebreak':
                boundaries: List[int] = []
                for index, block in enumerate(blocks):
                    if block.tag != qn('w:p'):
                        continue
                    if self._has_page_break(block) and index != 0:
                        boundaries.append(index)
                    if self._has_section_break(block) and index + 1 < total:
                        boundaries.append(index + 1)
                segments = self._segments_by_boundaries(total, boundaries)
                if len(segments) <= 1:
                    raise ValueError('未检测到分页符/分节符，无法按分页拆分')
            elif mode == 'heading':
                try:
                    id_to_name = {s.style_id: s.name for s in doc.styles}
                except Exception:
                    id_to_name = {}
                boundaries = [
                    index for index, block in enumerate(blocks)
                    if self._is_heading(block, id_to_name, heading_level)
                ]
                segments = self._segments_by_boundaries(total, boundaries)
                if len(segments) <= 1:
                    raise ValueError(f'未检测到 {heading_level} 级标题，无法按标题拆分')
            else:
                raise ValueError('未知的拆分模式')

            if len(segments) > 500:
                raise ValueError('拆分份数过多（>500），请调整参数')

            base_name = opts.get('outputName') or source.stem
            if base_name.lower().endswith('.docx'):
                base_name = base_name[:-5]

            exported: List[str] = []
            for part, segment in enumerate(iter_progress(segments, '正在拆分 Word'), start=1):
                dest = out_dir / f'{base_name}_part{part:03}.docx'
                dest = self._write_segment(source, dest, segment)
                exported.append(str(dest))

            return {
                'code': 0,
                'msg': f'拆分完成，共 {len(exported)} 个文件{extra_msg}',
                'files': exported,
                'outputDir': str(out_dir)
            }
        except Exception as exc:
            return {'code': -1, 'msg': f'拆分失败：{exc}'}

    def word_cut(self, options: Dict = None):
        '''Word 按真实页码切割：仅保留指定页码范围的内容，剔除其余，输出单个文件（保留格式）。

        模式（mode）：
          - range  : 保留 startPage~endPage
          - custom : 保留 pageSpec 表达式（如 "1-3,5,8"）所列页码
        '''
        try:
            opts = options if isinstance(options, dict) else {}
            source = self._ensure_word_file(opts.get('filePath', ''))
            mode = str(opts.get('mode') or 'range').lower()

            spans, total_pages = self._page_spans(source)

            if mode == 'range':
                start = max(1, int(opts.get('startPage') or 1))
                end = min(total_pages, int(opts.get('endPage') or start))
                if start > end:
                    raise ValueError('开始页不能大于结束页')
                target_pages = set(range(start, end + 1))
                suffix = 'range'
            else:
                pages = self._parse_page_spec(str(opts.get('pageSpec') or ''), total_pages)
                if not pages:
                    raise ValueError('请设置有效的页码')
                target_pages = set(pages)
                suffix = 'custom'

            keep = self._indices_for_pages(spans, target_pages)
            if not keep:
                raise ValueError('指定页码范围内没有内容')

            output_path = opts.get('outputPath') or self._compose_output_path(
                opts.get('outputDir', ''), opts.get('outputName', ''), source, suffix
            )
            dest = self._resolve_output_path(source, output_path, suffix)
            dest = self._write_segment(source, dest, keep)

            return {
                'code': 0,
                'msg': f'已保留 {len(target_pages)} 页内容（原文档共 {total_pages} 页）',
                'output': str(dest),
                'totalPages': total_pages
            }
        except Exception as exc:
            return {'code': -1, 'msg': f'切割失败：{exc}'}

    def word_merge(self, options: Dict = None):
        '''多个 Word(.docx) 按指定顺序合并为一个文档（保留各文档格式）。'''
        try:
            opts = options if isinstance(options, dict) else {}
            files = opts.get('files', [])
            if not isinstance(files, Iterable):
                raise ValueError('参数格式错误')
            page_break = bool(opts.get('pageBreak', True))

            paths: List[Path] = []
            for item in files:
                path = item.get('path', '') if isinstance(item, dict) else str(item)
                paths.append(self._ensure_word_file(path))

            if len(paths) < 2:
                raise ValueError('请至少选择 2 个 Word 文件')

            output_path = opts.get('outputPath') or self._compose_output_path(
                opts.get('outputDir', ''), opts.get('outputName', ''), paths[0], 'merged'
            )
            dest = self._resolve_output_path(paths[0], output_path, 'merged')

            master = Document(str(paths[0]))
            composer = Composer(master)
            for extra in iter_progress(paths[1:], '正在合并 Word'):
                if page_break:
                    master.add_page_break()
                composer.append(Document(str(extra)))
            dest = write_output(dest, lambda target: composer.save(str(target)))

            return {
                'code': 0,
                'msg': f'合并成功，共 {len(paths)} 个文件',
                'output': str(dest)
            }
        except Exception as exc:
            return {'code': -1, 'msg': f'合并失败：{exc}'}
