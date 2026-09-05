from __future__ import annotations

import json
import shutil
import sqlite3
import tempfile
import time
import unittest
import zipfile
from pathlib import Path
from unittest import mock

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from openpyxl import Workbook, load_workbook
from PIL import Image

from api.api import API
from api.core.context import run_process
from api.core.database import migrate_database
from api.core.store import StateStore
from api.operations import OPERATIONS
from api.pdf import PDF
from api.webauto import WebAutoTool
from api.word import WordTool
from api.workflow import WorkflowMixin
from pyapp.config.config import Config


class RefactorIntegrityTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.patch = mock.patch.object(Config, 'appDataDir', str(self.root / 'data'))
        self.patch.start()
        Path(Config.appDataDir).mkdir()
        self.api = API()

    def tearDown(self):
        self.api.task_shutdown()
        self.api.workflow_stop()
        self.patch.stop()
        self.temp.cleanup()

    def test_catalog_names_inputs_and_legacy_history_migration(self):
        for operation in OPERATIONS.values():
            self.assertNotIn('\ufffd', operation.label)
            for field in operation.fields:
                self.assertNotIn('\ufffd', field['label'])
        self.assertIn('files', {field['name'] for field in OPERATIONS['image_batch_compress'].fields})
        records = [{'id': 'old-1', 'method': 'image_rotate_flip', 'status': 'running', 'startedAt': 1700000000000, 'args': [{'files': ['original.png']}]}]
        self.assertEqual(self.api.task_import_legacy({'tasks': records})['imported'], 1)
        self.assertEqual(self.api.task_import_legacy({'tasks': records})['imported'], 0)
        self.assertEqual(self.api.task_get('old-1')['task']['status'], 'interrupted')
        self.assertEqual(self.api.task_get('old-1')['task']['createdAt'], 1700000000)

    def test_copy_thousand_files_including_empty_file(self):
        source, target = self.root / 'inputs', self.root / 'outputs'
        source.mkdir()
        for index in range(1000):
            (source / f'{index}.txt').write_bytes(b'' if index == 0 else str(index).encode())
        target.mkdir()
        (target / '1.txt').write_text('keep original', encoding='utf-8')
        result = self.api.file_batch_copy({'sourceDir': str(source), 'targetDir': str(target), 'conflictPolicy': 'overwrite'})
        self.assertEqual(result['code'], 0, result)
        self.assertEqual(result['copied'], 1000)
        self.assertEqual(len(result['outputAssets']), 1000)
        self.assertEqual((target / '0.txt').read_bytes(), b'')
        self.assertEqual((target / '1.txt').read_text(), 'keep original')

    def test_directory_retry_uses_only_failed_original_inputs(self):
        source, target = self.root / 'inputs', self.root / 'outputs'
        source.mkdir()
        # Exercise a noncanonical directory spelling on every platform, as well
        # as native /var aliases and Windows short temp paths on CI.
        source = source / '..' / 'inputs'
        for name in ('good.txt', 'blocked.txt'):
            (source / name).write_text(name, encoding='utf-8')
        original_copy = shutil.copy2

        def copy_with_failure(src, dst, *args, **kwargs):
            if Path(src).name == 'blocked.txt':
                raise PermissionError('file is busy')
            return original_copy(src, dst, *args, **kwargs)

        def wait_task(identity):
            deadline = time.monotonic() + 20
            while time.monotonic() < deadline:
                task = self.api.task_get(identity)['task']
                if task['status'] not in {'queued', 'running', 'canceling'}:
                    return task
                time.sleep(0.05)
            self.fail('task did not finish')

        with mock.patch('api.file.shutil.copy2', side_effect=copy_with_failure):
            submitted = self.api.task_submit({'method': 'file_batch_copy', 'args': [{'sourceDir': str(source), 'targetDir': str(target), 'conflictPolicy': 'overwrite'}]})
            partial = wait_task(submitted['taskId'])
        self.assertEqual(partial['status'], 'partial', partial)
        self.assertEqual([item['status'] for item in partial['itemResults']].count('success'), 1)
        (source / 'new.txt').write_text('not part of original batch', encoding='utf-8')
        retry = wait_task(self.api.task_retry(partial['id'])['taskId'])
        self.assertEqual(retry['status'], 'success', retry)
        self.assertEqual(retry['args'][0]['_retryInputs'], [str(source / 'blocked.txt')])
        self.assertEqual(sorted(path.name for path in target.iterdir()), ['blocked.txt', 'good.txt'])

    def test_schema_upgrade_backup_rollback_and_repeat(self):
        database = self.root / 'state.sqlite3'
        connection = sqlite3.connect(database)
        connection.execute('CREATE TABLE original (value TEXT)')
        connection.execute("INSERT INTO original VALUES ('preserve')")
        connection.commit()
        connection.close()

        def fail_upgrade(connection):
            connection.execute('CREATE TABLE incomplete (value TEXT)')
            raise ValueError('injected failure')

        with self.assertRaisesRegex(ValueError, 'injected'):
            migrate_database(database, {1: fail_upgrade})
        with sqlite3.connect(database) as connection:
            self.assertEqual(connection.execute('PRAGMA user_version').fetchone()[0], 0)
            self.assertEqual(connection.execute('SELECT value FROM original').fetchone()[0], 'preserve')
            self.assertIsNone(connection.execute("SELECT name FROM sqlite_master WHERE name='incomplete'").fetchone())
        connection.close()
        backup = migrate_database(database, {1: lambda db: db.execute('CREATE TABLE completed (id INTEGER)')})
        with sqlite3.connect(backup) as connection:
            self.assertEqual(connection.execute('PRAGMA user_version').fetchone()[0], 0)
            self.assertEqual(connection.execute('SELECT value FROM original').fetchone()[0], 'preserve')
        connection.close()
        self.assertIsNone(migrate_database(database, {1: lambda db: self.fail('migration repeated')}))

    def test_restore_legacy_backup_reimports_tasks_without_losing_other_namespaces(self):
        data = Path(Config.appDataDir)
        legacy = data / 'tasks/history.json'
        legacy.parent.mkdir()
        legacy.write_text(json.dumps({'tasks': [{'id': 'restored', 'method': 'text_case_transform', 'status': 'success'}]}), encoding='utf-8')
        backup = self.api.maintenance_backup_create()
        self.assertEqual(backup['code'], 0, backup)
        store = StateStore(data)
        store.save('tasks', {'tasks': [{'id': 'newer'}]})
        store.save('webauto', {'collections': [{'id': 'other'}]})
        staged = self.api.maintenance_backup_restore({'filePath': backup['output']})
        self.assertEqual(staged['code'], 0, staged)
        restored = self.api.maintenance_apply_pending_restore()
        self.assertEqual(restored['code'], 0, restored)
        self.assertEqual(store.load('tasks', legacy, {})['tasks'][0]['id'], 'restored')
        self.assertEqual(store.load('webauto', data / 'absent.json', {})['collections'][0]['id'], 'other')

    def test_excel_groups_keep_numeric_and_text_values_separate(self):
        source = self.root / 'types.xlsx'
        wb = Workbook()
        wb.active.append(['value'])
        wb.active.append([1])
        wb.active.append(['1'])
        wb.active.append(['001'])
        wb.save(source)
        wb.close()
        result = self.api.excel_process({'filePath': str(source), 'groupBy': 'value', 'outputDir': str(self.root / 'groups')})
        self.assertEqual(result['code'], 0, result)
        self.assertEqual(result['summary']['groupCount'], 3)

    def test_workflow_retry_reuses_successful_steps(self):
        class Workflow(WorkflowMixin):
            calls = []
            fail = True

            def text_case_transform(self, options=None):
                phase = options['phase']
                self.calls.append(phase)
                if phase == 'second' and self.fail:
                    return {'code': -1, 'msg': 'injected failure'}
                return {'code': 0, 'output': phase.upper(), 'outputAssets': []}

        workflow = Workflow()
        saved = workflow.workflow_save({'name': 'resume', 'steps': [
            {'id': 'one', 'method': 'text_case_transform', 'args': {'phase': 'first'}},
            {'id': 'two', 'method': 'text_case_transform', 'args': {'phase': 'second'}}]})
        identity = saved['workflow']['id']
        first = workflow.workflow_run({'id': identity})
        self.assertTrue(first['partial'], first)
        workflow.fail = False
        retry = workflow.workflow_run({'id': identity, '_resumeRunId': first['run']['id']})
        self.assertEqual(retry['code'], 0, retry)
        self.assertEqual(workflow.calls, ['first', 'second', 'second'])
        self.assertTrue(retry['run']['steps'][0]['reused'])

    def test_classification_undo_preserves_modified_output(self):
        source, target = self.root / 'inputs', self.root / 'organized'
        source.mkdir()
        (source / 'a.txt').write_text('alpha', encoding='utf-8')
        (source / 'b.txt').write_text('beta', encoding='utf-8')
        result = self.api.file_auto_classify({'directory': str(source), 'targetDir': str(target), 'operation': 'move'})
        self.assertEqual(result['code'], 0, result)
        changed = Path(next(item['to'] for item in result['operations'] if Path(item['from']).name == 'b.txt'))
        changed.write_text('edited after classification', encoding='utf-8')
        undone = self.api.file_classify_undo({'directory': str(source), 'transactionId': result['transactionId']})
        self.assertEqual(undone['restoredCount'], 1, undone)
        self.assertEqual((source / 'a.txt').read_text(), 'alpha')
        self.assertEqual(changed.read_text(), 'edited after classification')
        repeated = self.api.file_classify_undo({'directory': str(source), 'transactionId': result['transactionId']})
        self.assertEqual(repeated['restoredCount'], 0)

    def test_image_rename_undo_and_conflict(self):
        source = self.root / 'image.png'
        Image.new('RGBA', (40, 30), (1, 2, 3, 80)).save(source)
        original = source.read_bytes()
        result = self.api.image_batch_rename({'files': [str(source)], 'prefix': 'renamed_', 'dryRun': False})
        self.assertEqual(result['code'], 0, result)
        self.assertFalse(source.exists())
        undone = self.api.image_batch_rename_undo({'transactionId': result['transactionId']})
        self.assertEqual(undone['code'], 0, undone)
        self.assertEqual(source.read_bytes(), original)
        self.assertEqual(self.api.image_batch_rename_undo({'transactionId': result['transactionId']})['files'], [])

    def test_archive_never_overwrites_and_preserves_empty_files(self):
        source = self.root / 'source.txt'
        source.write_bytes(b'')
        archive = self.api.file_compress({'items': [str(source)], 'outputDir': str(self.root), 'archiveName': 'data'})
        self.assertEqual(archive['code'], 0, archive)
        target = self.root / 'unpack'
        target.mkdir()
        (target / source.name).write_text('existing', encoding='utf-8')
        result = self.api.file_decompress({'archiveFile': archive['file'], 'targetDir': str(target)})
        self.assertEqual(result['code'], 0, result)
        self.assertEqual((target / source.name).read_text(), 'existing')
        self.assertEqual(Path(result['files'][0]).read_bytes(), b'')

    def test_word_table_rows_sections_and_ambiguous_page_boundary(self):
        source = self.root / 'table.docx'
        document = Document()
        document.add_paragraph('first')
        table = document.add_table(rows=3, cols=1)
        for index, row in enumerate(table.rows):
            row.cells[0].text = f'row {index}'
        document.add_section(WD_SECTION.NEW_PAGE)
        document.add_paragraph('last')
        document.save(source)
        preview = self.api.word_preview({'filePath': str(source)})
        self.assertEqual(preview['sectionCount'], 2, preview)
        self.assertEqual([item['text'] for item in preview['blocks'] if item['kind'] == 'tableRow'], ['row 0', 'row 1', 'row 2'])
        dest = WordTool()._write_segment(source, self.root / 'segment.docx', [2])
        self.assertEqual(Document(dest).tables[0].cell(0, 0).text, 'row 1')
        with self.assertRaisesRegex(ValueError, '跨越'):
            WordTool()._indices_for_pages([(1, 1), (1, 2)], {1})

    def test_word_keeps_selected_section_layout_and_inherited_headers(self):
        source = self.root / 'sections.docx'
        document = Document()
        document.sections[0].header.paragraphs[0].text = 'Original header'
        document.add_paragraph('first section')
        document.add_section(WD_SECTION.NEW_PAGE)
        document.sections[1].left_margin = Inches(2)
        document.add_paragraph('second section')
        document.add_section(WD_SECTION.NEW_PAGE)
        document.sections[2].left_margin = Inches(0.5)
        document.add_paragraph('third section')
        document.save(source)
        result = WordTool()._write_segment(source, self.root / 'selected.docx', [2, 4])
        selected = Document(result)
        self.assertEqual(len(selected.sections), 2)
        self.assertEqual([section.left_margin for section in selected.sections], [Inches(2), Inches(0.5)])
        self.assertEqual([section.header.paragraphs[0].text for section in selected.sections], ['Original header'] * 2)

    def test_word_native_page_cut_has_only_selected_content(self):
        tool = WordTool()
        try:
            tool._locate_soffice()
        except RuntimeError:
            self.skipTest('LibreOffice is not installed on this runner')
        source = self.root / 'pages.docx'
        document = Document()
        document.add_paragraph('FIRST_PAGE_ONLY')
        document.add_page_break()
        document.add_paragraph('SECOND_PAGE_ONLY')
        document.save(source)
        result = self.api.word_cut({'filePath': str(source), 'startPage': 2, 'endPage': 2, 'outputDir': str(self.root / 'out')})
        self.assertEqual(result['code'], 0, result)
        text = '\n'.join(paragraph.text for paragraph in Document(result['output']).paragraphs)
        self.assertIn('SECOND_PAGE_ONLY', text)
        self.assertNotIn('FIRST_PAGE_ONLY', text)

    def test_excel_rejects_unsafe_formula_before_publishing(self):
        source = self.root / 'formula.xlsx'
        wb = Workbook()
        wb.active.append(['number', 'formula'])
        wb.active.append([1, '=$A$2+1'])
        wb.save(source)
        wb.close()
        result = self.api.excel_process({'filePath': str(source), 'exportCombined': True, 'outputDir': str(self.root / 'out'), 'formulaPolicy': 'preserve'})
        self.assertNotEqual(result['code'], 0)
        self.assertEqual(list((self.root / 'out').glob('*.xlsx')), [])

    def test_document_index_locations_and_ocr_increment(self):
        source = self.root / 'pages.pdf'
        document = fitz.open()
        document.new_page().insert_text((30, 50), 'first topic')
        document.new_page().insert_text((30, 50), 'needle topic')
        document.save(source)
        document.close()
        self.assertEqual(self.api.document_index_build({'files': [str(source)]})['indexed'], 1)
        hits = self.api.document_index_search({'query': 'needle'})['results']
        self.assertEqual(hits[0]['locations'][0]['page'], 2)
        source.write_bytes(source.read_bytes())
        self.assertTrue(self.api.document_index_search({'query': 'needle'})['results'][0]['stale'])

    def test_wal_backup_contains_committed_rows(self):
        database = Path(Config.appDataDir) / 'example.sqlite3'
        connection = sqlite3.connect(database)
        try:
            connection.execute('PRAGMA journal_mode=WAL')
            connection.execute('PRAGMA wal_autocheckpoint=0')
            connection.execute('CREATE TABLE example(value TEXT)')
            connection.execute("INSERT INTO example VALUES ('committed in WAL')")
            connection.commit()
            backup = self.api.maintenance_backup_create({'outputDir': str(self.root / 'backups')})
            self.assertEqual(backup['code'], 0, backup)
            with zipfile.ZipFile(backup['output']) as archive:
                snapshot = self.root / 'snapshot.sqlite3'
                snapshot.write_bytes(archive.read('data/example.sqlite3'))
            with sqlite3.connect(snapshot) as restored:
                self.assertEqual(restored.execute('SELECT value FROM example').fetchone()[0], 'committed in WAL')
            restored.close()
        finally:
            connection.close()

    def test_saved_collection_export_keeps_formula_text_and_old_results(self):
        web = WebAutoTool()
        web._wa_ensure()
        web._wa_save_result([{'value': '=SUM(1,2)'}], ['value'])
        first_id = web._wa_run['resultId']
        web._wa_run = web._wa_blank_run()
        web._wa_save_result([{'value': 'second collection'}], ['value'])
        result = self.api.webauto_export({'resultId': first_id, 'outputDir': str(self.root), 'fileName': 'first'})
        self.assertEqual(result['code'], 0, result)
        workbook = load_workbook(result['outputPath'])
        try:
            self.assertEqual(workbook.active['A2'].value, '=SUM(1,2)')
            self.assertEqual(workbook.active['A2'].data_type, 's')
        finally:
            workbook.close()

    @unittest.skipUnless(shutil.which('ffmpeg') and shutil.which('ffprobe'), 'FFmpeg toolchain is optional')
    def test_video_fractional_cut_has_valid_duration_and_stream(self):
        source = self.root / 'source.mp4'
        generated = run_process([shutil.which('ffmpeg'), '-v', 'error', '-f', 'lavfi', '-i', 'color=c=blue:s=128x96:r=25:d=2', '-c:v', 'libx264', '-y', str(source)])
        self.assertEqual(generated.returncode, 0, generated.stderr)
        result = self.api.video_cut({'filePath': str(source), 'start': 0.4, 'end': 1.2, 'outputDir': str(self.root)})
        self.assertEqual(result['code'], 0, result)
        inspected = self.api.video_inspect({'filePath': result['file']})
        self.assertAlmostEqual(float(inspected['format']['duration']), 0.8, delta=0.12)
        self.assertEqual(inspected['streams'][0]['codec_name'], 'h264')

    @unittest.skipUnless(shutil.which('ffmpeg') and shutil.which('ffprobe'), 'FFmpeg toolchain is optional')
    def test_video_concat_checks_and_normalizes_different_dimensions(self):
        files = []
        for index, size in enumerate(['128x96', '160x120']):
            source = self.root / f'clip{index}.mp4'
            generated = run_process([shutil.which('ffmpeg'), '-v', 'error', '-f', 'lavfi', '-i', f'color=c=blue:s={size}:r=25:d=0.8', '-c:v', 'libx264', '-y', str(source)])
            self.assertEqual(generated.returncode, 0, generated.stderr)
            files.append(str(source))
        self.assertFalse(self.api.video_concat_preview({'files': files})['compatible'])
        rejected = self.api.video_concat({'files': files, 'outputDir': str(self.root / 'out')})
        self.assertNotEqual(rejected['code'], 0)
        result = self.api.video_concat({'files': files, 'outputDir': str(self.root / 'out'), 'reencode': True})
        self.assertEqual(result['code'], 0, result)
        inspected = self.api.video_inspect({'filePath': result['file']})
        self.assertAlmostEqual(float(inspected['format']['duration']), 1.6, delta=0.15)
        self.assertEqual((inspected['streams'][0]['width'], inspected['streams'][0]['height']), (128, 96))

    def test_pdf_page_failure_retries_only_the_failed_page(self):
        source = self.root / 'partial.pdf'
        document = fitz.open()
        document.new_page(width=64, height=80)
        document.new_page(width=65, height=80)
        document.save(source)
        document.close()
        output_dir = self.root / 'images'
        original = PDF._pil_from_pixmap

        def fail_second_page(tool, pixmap):
            if pixmap.width == 65:
                raise ValueError('injected page decode failure')
            return original(tool, pixmap)

        def wait_task(identity):
            deadline = time.monotonic() + 20
            while time.monotonic() < deadline:
                task = self.api.task_get(identity)['task']
                if task['status'] not in {'queued', 'running', 'canceling'}:
                    return task
                time.sleep(0.05)
            self.fail('PDF task did not finish')

        with mock.patch.object(PDF, '_pil_from_pixmap', fail_second_page), mock.patch('api.tasks.run_in_worker', side_effect=lambda method, args, context: getattr(self.api, method)(*args)):
            submission = self.api.task_submit({'method': 'pdf_convert_to_images', 'args': [{'filePath': str(source), 'dpi': 72, 'outputDir': str(output_dir)}]})
            partial = wait_task(submission['taskId'])
        self.assertEqual(partial['status'], 'partial', partial)
        self.assertEqual(len(partial['outputs']), 1)
        retry = wait_task(self.api.task_retry(partial['id'])['taskId'])
        self.assertEqual(retry['status'], 'success', retry)
        self.assertEqual(retry['args'][0]['_retryPageNumbers'], [2])
        self.assertEqual(len(list(output_dir.glob('*.png'))), 2)


if __name__ == '__main__':
    unittest.main()
